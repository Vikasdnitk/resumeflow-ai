import io
import json
import os
import re
import sqlite3
import tempfile
import uuid
from collections import Counter
from datetime import datetime

from flask import Flask, Response, jsonify, render_template, request, session

try:
    import docx
except Exception:
    docx = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.utils import simpleSplit
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None
    colors = None
    simpleSplit = None
    canvas = None


app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-change-me")
CACHE_DIR = os.path.join(tempfile.gettempdir(), "ats_resume_cache")
os.makedirs(CACHE_DIR, exist_ok=True)
DB_PATH = os.path.join(os.path.dirname(__file__), "ats_resume.db")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")


def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db_connection()
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS generations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT NOT NULL,
            strict_mode INTEGER NOT NULL,
            match_score INTEGER NOT NULL,
            jd_preview TEXT NOT NULL,
            tailored_preview TEXT NOT NULL,
            payload_id TEXT NOT NULL
        )
        """
    )
    conn.commit()
    conn.close()


def save_generation_record(strict_mode, match_score, job_description, tailored_resume, payload_id):
    jd_preview = " ".join(job_description.split())[:240]
    tailored_preview = " ".join(tailored_resume.split())[:320]
    conn = get_db_connection()
    conn.execute(
        """
        INSERT INTO generations (created_at, strict_mode, match_score, jd_preview, tailored_preview, payload_id)
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            datetime.utcnow().isoformat(timespec="seconds"),
            1 if strict_mode else 0,
            int(match_score),
            jd_preview,
            tailored_preview,
            payload_id,
        ),
    )
    conn.commit()
    conn.close()


def fetch_recent_generations(limit=5):
    conn = get_db_connection()
    rows = conn.execute(
        """
        SELECT id, created_at, strict_mode, match_score, jd_preview
        FROM generations
        ORDER BY id DESC
        LIMIT ?
        """,
        (limit,),
    ).fetchall()
    conn.close()
    return [dict(row) for row in rows]


init_db()


STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "has",
    "he",
    "in",
    "is",
    "it",
    "its",
    "of",
    "on",
    "that",
    "the",
    "to",
    "was",
    "were",
    "will",
    "with",
    "you",
    "your",
    "or",
    "we",
    "our",
    "they",
    "their",
    "this",
    "those",
    "these",
    "job",
    "role",
    "candidate",
    "experience",
    "skills",
    "responsibilities",
    "requirements",
}


def extract_text_from_file(file_storage):
    filename = (file_storage.filename or "").lower()
    content = file_storage.read()

    if filename.endswith(".txt"):
        return content.decode("utf-8", errors="ignore")

    if filename.endswith(".pdf"):
        if PdfReader is None:
            raise ValueError("PDF support is not installed. Install pypdf.")
        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    if filename.endswith(".docx"):
        if docx is None:
            raise ValueError("DOCX support is not installed. Install python-docx.")
        document = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in document.paragraphs)

    raise ValueError("Unsupported file format. Use .txt, .pdf, or .docx")


def tokenize(text):
    words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-]{1,}", text.lower())
    return [w for w in words if w not in STOPWORDS and len(w) > 2]


def extract_keywords(job_description, limit=30):
    tokens = tokenize(job_description)
    freq = Counter(tokens)
    ranked = [token for token, _ in freq.most_common(limit * 2)]

    keywords = []
    seen = set()
    for token in ranked:
        if token not in seen:
            keywords.append(token)
            seen.add(token)
        if len(keywords) >= limit:
            break
    return keywords


def split_resume_sections(resume_text):
    sections = {
        "summary": "",
        "experience": "",
        "projects": "",
        "education": "",
        "skills": "",
    }

    lines = [line.strip() for line in resume_text.splitlines() if line.strip()]
    current = "summary"

    section_map = {
        "summary": "summary",
        "professional summary": "summary",
        "experience": "experience",
        "work experience": "experience",
        "projects": "projects",
        "education": "education",
        "skills": "skills",
        "technical skills": "skills",
    }

    for line in lines:
        normalized = line.lower().strip(":")
        if normalized in section_map:
            current = section_map[normalized]
            continue
        sections[current] += (line + "\n")

    for key, value in sections.items():
        sections[key] = value.strip()

    return sections


def split_any_resume_sections(resume_text):
    sections = {
        "summary": [],
        "experience": [],
        "projects": [],
        "education": [],
        "skills": [],
        "other": [],
    }
    current = "summary"
    section_map = {
        "summary": "summary",
        "professional summary": "summary",
        "profile": "summary",
        "experience": "experience",
        "work experience": "experience",
        "employment": "experience",
        "projects": "projects",
        "project": "projects",
        "education": "education",
        "skills": "skills",
        "technical skills": "skills",
        "certifications": "other",
        "achievements": "other",
    }

    for raw in resume_text.splitlines():
        line = raw.strip()
        if not line:
            continue
        normalized = line.lower().strip(":")
        if normalized in section_map:
            current = section_map[normalized]
            continue
        sections[current].append(line)
    return sections


def parse_resume_blocks(full_text):
    blocks = []
    current_title = "RESUME"
    current_lines = []
    known = {
        "summary",
        "professional summary",
        "profile",
        "experience",
        "work experience",
        "employment",
        "projects",
        "project",
        "education",
        "skills",
        "technical skills",
        "certifications",
        "achievements",
        "contact",
        "personal details",
    }

    for raw in full_text.splitlines():
        line = raw.strip()
        if not line:
            continue
        normalized = line.lower().strip(":")
        is_heading = normalized in known or (line.isupper() and len(line.split()) <= 4)
        if is_heading:
            if current_lines:
                blocks.append((current_title, current_lines))
            current_title = line.upper()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        blocks.append((current_title, current_lines))
    return blocks


def infer_candidate_name(resume_text):
    for line in resume_text.splitlines()[:6]:
        cleaned = line.strip()
        if cleaned and len(cleaned.split()) <= 4 and not any(char.isdigit() for char in cleaned):
            if not re.search(r"resume|curriculum|vitae|summary|experience", cleaned, re.I):
                return cleaned.title()
    return "Candidate Name"


def infer_contact_line(resume_text):
    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", resume_text)
    phone = re.search(r"(\+?\d[\d\-\s()]{8,}\d)", resume_text)
    linkedin = re.search(r"(linkedin\.com/\S+)", resume_text, re.I)
    parts = []
    if email:
        parts.append(email.group(0))
    if phone:
        parts.append(re.sub(r"\s+", " ", phone.group(0).strip()))
    if linkedin:
        parts.append(linkedin.group(1))
    return " | ".join(parts) if parts else "Email | Phone | LinkedIn"


def to_bullets(text, limit=20):
    items = []
    for raw_line in text.splitlines():
        line = raw_line.strip().lstrip("-*").strip()
        if line:
            items.append(line)
    if not items:
        return []
    return items[:limit]


def optimize_existing_resume_text(resume_text, keywords):
    lines = resume_text.splitlines()
    if not lines:
        return resume_text

    heading_map = {
        "summary",
        "professional summary",
        "experience",
        "work experience",
        "projects",
        "education",
        "skills",
        "technical skills",
    }

    updated = []
    used_keywords = set()
    kw_index = 0

    for raw in lines:
        line = raw.rstrip()
        normalized = line.lower().strip().strip(":")

        if normalized in heading_map or not line.strip():
            updated.append(line)
            continue

        # Keep original content and only enrich a few bullets with JD terms.
        if (line.lstrip().startswith(("-", "*", "•")) or len(line.split()) > 7) and kw_index < len(keywords):
            additions = []
            for _ in range(2):
                if kw_index < len(keywords):
                    kw = keywords[kw_index]
                    kw_index += 1
                    if kw not in used_keywords and kw.lower() not in line.lower():
                        additions.append(kw)
                        used_keywords.add(kw)
            if additions:
                updated.append(f"{line} | Tools/Keywords: {', '.join(additions)}")
                continue

        updated.append(line)

    # Ensure skills section contains important missing keywords.
    text_joined = "\n".join(updated)
    if "skills" in text_joined.lower():
        missing = [kw for kw in keywords[:20] if kw.lower() not in text_joined.lower()]
        if missing:
            updated.append("")
            updated.append("Additional ATS Keywords")
            updated.append(", ".join(missing[:12]))

    return "\n".join(updated).strip()


def add_ats_keywords(existing_skills, keywords, limit=20):
    current = {s.strip().lower() for s in re.split(r",|\n|\|", existing_skills) if s.strip()}
    prioritized = [kw for kw in keywords if kw.lower() not in current]
    merged = [s for s in re.split(r",|\n|\|", existing_skills) if s.strip()]
    merged.extend(prioritized[:limit])
    return ", ".join(dict.fromkeys([m.strip() for m in merged if m.strip()]))


def build_tailored_resume(resume_text, job_description):
    sections = split_resume_sections(resume_text)
    keywords = extract_keywords(job_description)

    top_keywords = ", ".join(keywords[:12])

    summary = sections["summary"] or "Results-driven professional with strong ownership and collaboration skills."
    experience = sections["experience"] or "Delivered high-impact work with ownership and measurable outcomes."
    projects = sections["projects"] or "Built projects relevant to this role with clear technical impact."
    education = sections["education"] or "Add your most relevant education details."
    skills = add_ats_keywords(sections["skills"], keywords)

    tailored_summary = f"{summary} Core strengths include {top_keywords}."
    optimized_text = optimize_existing_resume_text(resume_text, keywords)

    structured = {
        "name": infer_candidate_name(resume_text),
        "contact": infer_contact_line(resume_text),
        "summary": tailored_summary,
        "experience": to_bullets(experience),
        "projects": to_bullets(projects),
        "skills": [s.strip() for s in skills.split(",") if s.strip()][:30],
        "education": to_bullets(education, limit=10),
    }

    if not structured["experience"]:
        structured["experience"] = [
            "Led cross-functional delivery and improved process efficiency.",
            "Built solutions aligned to role requirements and business goals.",
        ]

    if not structured["projects"]:
        structured["projects"] = [
            "Implemented role-relevant project using modern tools and best practices.",
        ]

    if not structured["education"]:
        structured["education"] = ["Bachelor's Degree - Institution Name"]

    lines = [
        structured["name"],
        structured["contact"],
        "",
        "PROFESSIONAL SUMMARY",
        structured["summary"],
        "",
        "EXPERIENCE",
    ]
    lines.extend([f"- {item}" for item in structured["experience"]])
    lines.extend(["", "PROJECTS"])
    lines.extend([f"- {item}" for item in structured["projects"]])
    lines.extend(["", "SKILLS", ", ".join(structured["skills"]), "", "EDUCATION"])
    lines.extend([f"- {item}" for item in structured["education"]])

    structured_text = "\n".join(lines).strip()
    final_text = optimized_text if len(optimized_text) > 80 else structured_text
    return final_text, keywords, structured


def build_strict_resume_text(resume_text, keywords):
    cleaned = resume_text.strip()
    if not cleaned:
        cleaned = ""
    existing_tokens = set(tokenize(cleaned))
    missing = [kw for kw in keywords if kw.lower() not in existing_tokens]
    keyword_block = ", ".join(missing[:30]) if missing else "No major missing keywords detected."
    return (
        f"{cleaned}\n\n"
        "ATS KEYWORDS (FROM JOB DESCRIPTION)\n"
        f"{keyword_block}\n"
    ).strip()


def estimate_match_score(resume_text, keywords):
    resume_tokens = set(tokenize(resume_text))
    if not keywords:
        return 0, []
    matched = [kw for kw in keywords if kw.lower() in resume_tokens]
    score = int((len(matched) / len(keywords)) * 100)
    missing = [kw for kw in keywords if kw not in matched][:15]
    return score, missing


def build_docx_bytes(structured_resume):
    if docx is None:
        raise ValueError("DOCX export requires python-docx.")
    document = docx.Document()
    document.add_heading(structured_resume["name"], level=0)
    document.add_paragraph(structured_resume["contact"])
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(structured_resume["summary"])

    document.add_heading("Experience", level=1)
    for item in structured_resume["experience"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Projects", level=1)
    for item in structured_resume["projects"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Skills", level=1)
    document.add_paragraph(", ".join(structured_resume["skills"]))

    document.add_heading("Education", level=1)
    for item in structured_resume["education"]:
        document.add_paragraph(item, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(structured_resume):
    if canvas is None or A4 is None:
        raise ValueError("PDF export requires reportlab.")
    width, height = A4
    left_margin = 50
    y = height - 50
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    def draw_line(text, size=10, bold=False):
        nonlocal y
        font_name = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font_name, size)
        wrapped = simpleSplit(text, font_name, size, width - (left_margin * 2))
        for line in wrapped:
            if y < 60:
                c.showPage()
                y = height - 50
                c.setFont(font_name, size)
            c.drawString(left_margin, y, line)
            y -= size + 4

    draw_line(structured_resume["name"], size=16, bold=True)
    draw_line(structured_resume["contact"], size=10)
    y -= 6

    for section, content in [
        ("PROFESSIONAL SUMMARY", [structured_resume["summary"]]),
        ("EXPERIENCE", [f"- {item}" for item in structured_resume["experience"]]),
        ("PROJECTS", [f"- {item}" for item in structured_resume["projects"]]),
        ("SKILLS", [", ".join(structured_resume["skills"])]),
        ("EDUCATION", [f"- {item}" for item in structured_resume["education"]]),
    ]:
        draw_line(section, size=11, bold=True)
        for line in content:
            draw_line(line, size=10)
        y -= 5

    c.save()
    return buffer.getvalue()


def draw_section_with_line(c, width, y, title):
    c.setFont("Helvetica-Bold", 11)
    c.drawString(50, y, title)
    y -= 6
    c.setLineWidth(0.8)
    c.line(50, y, width - 50, y)
    return y - 14


def build_pdf_template_1(structured_resume, full_text=""):
    if canvas is None or A4 is None:
        raise ValueError("PDF export requires reportlab.")
    width, height = A4
    y = height - 52
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    def draw_wrapped(text, size=10, left=50, max_width=None, gap=4):
        nonlocal y
        max_width = max_width or (width - 100)
        font_name = "Helvetica" if size <= 10 else "Helvetica-Bold"
        c.setFont(font_name, size)
        wrapped = simpleSplit(text, font_name, size, max_width)
        for line in wrapped:
            if y < 60:
                c.showPage()
                y = height - 52
                c.setFont(font_name, size)
            c.drawString(left, y, line)
            y -= size + gap

    blocks = parse_resume_blocks(full_text)

    c.setFillColor(colors.HexColor("#0B2239") if colors else (0, 0, 0))
    c.setFont("Times-Bold", 20)
    c.drawString(50, y, structured_resume["name"])
    y -= 22
    y = draw_section_with_line(c, width, y, "PERSONAL DETAILS")
    c.setFillColor(colors.black if colors else (0, 0, 0))
    draw_wrapped(structured_resume["contact"], size=10)

    for title, lines in blocks:
        y = draw_section_with_line(c, width, y, title)
        for line in lines:
            draw_wrapped(line, size=10)

    c.save()
    return buffer.getvalue()


def build_pdf_template_2(structured_resume, full_text=""):
    if canvas is None or A4 is None:
        raise ValueError("PDF export requires reportlab.")
    width, height = A4
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    left_x = 44
    mid_x = 208
    right_x = 228
    divider_x = 220
    top_y = height - 48
    bottom_y = 45
    left_w = mid_x - left_x - 6
    right_w = width - right_x - 50

    parsed = split_any_resume_sections(full_text)

    def build_column_lines(side):
        lines = []
        if side == "left":
            lines.append(("CONTACT", "Helvetica-Bold", 11, 14))
            for line in simpleSplit(structured_resume["contact"], "Helvetica", 9, left_w):
                lines.append((line, "Helvetica", 9, 11))
            lines.append(("", "Helvetica", 9, 8))
            lines.append(("SKILLS", "Helvetica-Bold", 11, 14))
            left_skills = parsed["skills"] if parsed["skills"] else structured_resume["skills"]
            for skill in left_skills:
                for line in simpleSplit(f"- {skill}", "Helvetica", 9, left_w):
                    lines.append((line, "Helvetica", 9, 10))
            lines.append(("", "Helvetica", 9, 8))
            lines.append(("EDUCATION", "Helvetica-Bold", 11, 14))
            left_edu = parsed["education"] if parsed["education"] else structured_resume["education"]
            for edu in left_edu:
                for line in simpleSplit(f"- {edu}", "Helvetica", 9, left_w):
                    lines.append((line, "Helvetica", 9, 10))
            if parsed["other"]:
                lines.append(("", "Helvetica", 9, 8))
                lines.append(("ADDITIONAL", "Helvetica-Bold", 11, 14))
                for item in parsed["other"]:
                    for line in simpleSplit(f"- {item}", "Helvetica", 9, left_w):
                        lines.append((line, "Helvetica", 9, 10))
            return lines

        lines.append((structured_resume["name"], "Helvetica-Bold", 17, 20))
        lines.append(("Targeted ATS Resume", "Helvetica", 9, 14))
        lines.append(("PROFESSIONAL SUMMARY", "Helvetica-Bold", 11, 12))
        summary_text = " ".join(parsed["summary"]) if parsed["summary"] else structured_resume["summary"]
        for line in simpleSplit(summary_text, "Helvetica", 9.5, right_w):
            lines.append((line, "Helvetica", 9.5, 12))
        lines.append(("", "Helvetica", 9, 7))
        lines.append(("EXPERIENCE", "Helvetica-Bold", 11, 12))
        right_exp = parsed["experience"] if parsed["experience"] else structured_resume["experience"]
        for item in right_exp:
            for line in simpleSplit(f"- {item}", "Helvetica", 9.5, right_w):
                lines.append((line, "Helvetica", 9.5, 12))
        lines.append(("", "Helvetica", 9, 7))
        lines.append(("PROJECTS", "Helvetica-Bold", 11, 12))
        right_proj = parsed["projects"] if parsed["projects"] else structured_resume["projects"]
        for item in right_proj:
            for line in simpleSplit(f"- {item}", "Helvetica", 9.5, right_w):
                lines.append((line, "Helvetica", 9.5, 12))
        if parsed["other"]:
            lines.append(("", "Helvetica", 9, 7))
            lines.append(("ADDITIONAL", "Helvetica-Bold", 11, 12))
            for item in parsed["other"]:
                for line in simpleSplit(f"- {item}", "Helvetica", 9.5, right_w):
                    lines.append((line, "Helvetica", 9.5, 12))
        return lines

    def paginate(lines):
        pages = []
        current = []
        used = 0
        capacity = top_y - bottom_y
        for item in lines:
            step = item[3]
            if used + step > capacity and current:
                pages.append(current)
                current = []
                used = 0
            current.append(item)
            used += step
        if current:
            pages.append(current)
        return pages

    left_pages = paginate(build_column_lines("left"))
    right_pages = paginate(build_column_lines("right"))
    page_count = max(len(left_pages), len(right_pages))

    def draw_column(page_lines, x, is_right=False):
        y = top_y
        for text, font, size, step in page_lines:
            c.setFont(font, size)
            if is_right and text == "Targeted ATS Resume":
                c.setFillColorRGB(0.3, 0.34, 0.37)
                c.drawString(x, y, text)
                c.setFillColorRGB(0, 0, 0)
            else:
                c.drawString(x, y, text)
            y -= step

    for page_index in range(page_count):
        c.setStrokeColor(colors.HexColor("#0A4D68") if colors else (0.2, 0.35, 0.4))
        c.setLineWidth(2.2)
        c.line(44, height - 34, width - 44, height - 34)
        c.setStrokeColor(colors.HexColor("#B8C4CC") if colors else (0.75, 0.77, 0.78))
        c.setLineWidth(1)
        c.line(divider_x, 34, divider_x, height - 34)

        draw_column(left_pages[page_index] if page_index < len(left_pages) else [], left_x, is_right=False)
        draw_column(right_pages[page_index] if page_index < len(right_pages) else [], right_x, is_right=True)

        if page_index < page_count - 1:
            c.showPage()

    c.save()
    return buffer.getvalue()


def build_pdf_template_3(structured_resume, full_text=""):
    if canvas is None or A4 is None:
        raise ValueError("PDF export requires reportlab.")
    width, height = A4
    left = 48
    right = width - 48
    y = height - 52
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    def ensure_space(space_needed=40):
        nonlocal y
        if y < space_needed:
            c.showPage()
            y = height - 52
            draw_header(False)

    def draw_header(first_page=True):
        nonlocal y
        c.setFillColor(colors.HexColor("#0E2A47") if colors else (0, 0, 0))
        c.setFont("Times-Bold", 19 if first_page else 16)
        c.drawString(left, y, structured_resume["name"])
        y -= 18
        c.setFillColor(colors.HexColor("#3E4C59") if colors else (0, 0, 0))
        c.setFont("Helvetica", 9.5)
        for ln in simpleSplit(structured_resume["contact"], "Helvetica", 9.5, right - left):
            c.drawString(left, y, ln)
            y -= 12
        c.setStrokeColor(colors.HexColor("#C2CCD6") if colors else (0.8, 0.8, 0.8))
        c.setLineWidth(1)
        c.line(left, y, right, y)
        y -= 14
        c.setFillColor(colors.black if colors else (0, 0, 0))

    def section(title, lines):
        nonlocal y
        ensure_space(80)
        c.setFillColor(colors.HexColor("#0A4D68") if colors else (0, 0, 0))
        c.setFont("Helvetica-Bold", 11)
        c.drawString(left, y, title)
        y -= 12
        c.setFillColor(colors.black if colors else (0, 0, 0))
        c.setFont("Helvetica", 10)
        for line in lines:
            wrapped = simpleSplit(line, "Helvetica", 10, right - left)
            for wline in wrapped:
                ensure_space(55)
                c.drawString(left, y, wline)
                y -= 12
        y -= 7

    draw_header(True)
    section("PROFESSIONAL SUMMARY", [structured_resume.get("summary", "")])
    section("EXPERIENCE", [f"- {x}" for x in structured_resume.get("experience", [])])
    section("PROJECTS", [f"- {x}" for x in structured_resume.get("projects", [])])
    section("SKILLS", [", ".join(structured_resume.get("skills", []))])
    section("EDUCATION", [f"- {x}" for x in structured_resume.get("education", [])])

    c.save()
    return buffer.getvalue()


def build_pdf_by_template(structured_resume, template_id, full_text=""):
    if template_id == "t1":
        return build_pdf_template_1(structured_resume, full_text=full_text)
    if template_id == "t2":
        return build_pdf_template_2(structured_resume, full_text=full_text)
    return build_pdf_template_3(structured_resume, full_text=full_text)


def get_latest_structured_resume():
    payload_id = session.get("latest_resume_payload_id")
    if not payload_id:
        raise ValueError("Generate a resume first, then download it.")
    payload_path = os.path.join(CACHE_DIR, f"{payload_id}.json")
    if not os.path.exists(payload_path):
        raise ValueError("Generate a resume first, then download it.")
    with open(payload_path, "r", encoding="utf-8") as fh:
        return json.load(fh)


def save_resume_payload(tailored_resume, structured_resume):
    payload_id = uuid.uuid4().hex
    payload_path = os.path.join(CACHE_DIR, f"{payload_id}.json")
    payload = {
        "tailored_resume": tailored_resume,
        "structured_resume": structured_resume,
        "keywords": [],
        "missing_keywords": [],
        "match_score": 0,
    }
    with open(payload_path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh)
    return payload_id


def update_resume_payload(payload_id, keywords, missing_keywords, match_score):
    payload_path = os.path.join(CACHE_DIR, f"{payload_id}.json")
    if not os.path.exists(payload_path):
        return
    with open(payload_path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)
    payload["keywords"] = keywords
    payload["missing_keywords"] = missing_keywords
    payload["match_score"] = int(match_score)
    with open(payload_path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh)


def apply_user_modifications(resume_text, modification_text, keywords):
    updated = resume_text.strip()
    extra_keywords = extract_keywords(modification_text, limit=12)
    merged_keywords = list(dict.fromkeys(keywords + extra_keywords))

    if "one page" in modification_text.lower() or "short" in modification_text.lower():
        lines = [ln for ln in updated.splitlines() if ln.strip()]
        updated = "\n".join(lines[:70])

    updated = optimize_existing_resume_text(updated, merged_keywords)
    updated = (
        f"{updated}\n\n"
        "USER REQUESTED MODIFICATIONS\n"
        f"{modification_text.strip()}\n"
    ).strip()
    return updated, merged_keywords


def build_chat_context(payload):
    structured = payload.get("structured_resume", {})
    return {
        "resume_text": payload.get("tailored_resume", ""),
        "summary": structured.get("summary", ""),
        "experience": structured.get("experience", []),
        "projects": structured.get("projects", []),
        "skills": structured.get("skills", []),
        "education": structured.get("education", []),
    }


@app.route("/download/<filetype>", methods=["GET"])
def download_resume(filetype):
    try:
        payload = get_latest_structured_resume()
        structured_resume = payload.get("structured_resume", {})
        latest_tailored_resume = payload.get("tailored_resume", "")
        if filetype == "txt":
            content = latest_tailored_resume
            return Response(
                content,
                mimetype="text/plain",
                headers={"Content-Disposition": "attachment; filename=ats_resume.txt"},
            )

        if filetype == "docx":
            content = build_docx_bytes(structured_resume)
            return Response(
                content,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=ats_resume.docx"},
            )

        if filetype == "pdf":
            template_id = request.args.get("template", "t3").strip().lower()
            content = build_pdf_by_template(structured_resume, template_id, full_text=latest_tailored_resume)
            return Response(
                content,
                mimetype="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=ats_resume_{template_id}.pdf"},
            )

        return Response("Unsupported file type.", status=400)
    except Exception as exc:
        return Response(f"Download failed: {exc}", status=400)


@app.route("/", methods=["GET", "POST"])
def home():
    context = {
        "error": None,
        "tailored_resume": None,
        "match_score": None,
        "missing_keywords": [],
        "recent_generations": fetch_recent_generations(),
    }

    if request.method == "GET":
        session.pop("latest_resume_payload_id", None)
        return render_template("index.html", **context)

    if request.method == "POST":
        job_description = request.form.get("job_description", "").strip()
        resume_file = request.files.get("resume")
        strict_mode = request.form.get("strict_mode") == "on"

        if not job_description:
            context["error"] = "Please paste a job description."
            return render_template("index.html", **context)

        if not resume_file or not resume_file.filename:
            context["error"] = "Please upload your resume."
            return render_template("index.html", **context)

        try:
            resume_text = extract_text_from_file(resume_file)
            if strict_mode:
                keywords = extract_keywords(job_description)
                tailored_resume = build_strict_resume_text(resume_text, keywords)
                sections = split_resume_sections(resume_text)
                structured_resume = {
                    "name": infer_candidate_name(resume_text),
                    "contact": infer_contact_line(resume_text),
                    "summary": sections["summary"] or "",
                    "experience": to_bullets(sections["experience"], limit=50),
                    "projects": to_bullets(sections["projects"], limit=50),
                    "skills": [s.strip() for s in re.split(r",|\n|\|", sections["skills"]) if s.strip()],
                    "education": to_bullets(sections["education"], limit=30),
                }
                if not structured_resume["experience"]:
                    structured_resume["experience"] = ["Experience details are preserved in full text output."]
                if not structured_resume["projects"]:
                    structured_resume["projects"] = ["Project details are preserved in full text output."]
                if not structured_resume["education"]:
                    structured_resume["education"] = ["Education details are preserved in full text output."]
                if not structured_resume["skills"]:
                    structured_resume["skills"] = ["Skills are preserved in full text output."]
            else:
                tailored_resume, keywords, structured_resume = build_tailored_resume(resume_text, job_description)
            score, missing = estimate_match_score(tailored_resume, keywords)
            payload_id = save_resume_payload(tailored_resume, structured_resume)
            session["latest_resume_payload_id"] = payload_id
            update_resume_payload(payload_id, keywords, missing, score)
            save_generation_record(strict_mode, score, job_description, tailored_resume, payload_id)
            context["recent_generations"] = fetch_recent_generations()

            context.update(
                {
                    "tailored_resume": tailored_resume,
                    "match_score": score,
                    "missing_keywords": missing,
                    "download_ready": True,
                }
            )
        except Exception as exc:
            context["error"] = f"Could not process resume: {exc}"

    return render_template("index.html", **context)


@app.route("/load/<int:generation_id>", methods=["GET"])
def load_generation(generation_id):
    conn = get_db_connection()
    row = conn.execute(
        "SELECT payload_id FROM generations WHERE id = ?",
        (generation_id,),
    ).fetchone()
    conn.close()
    if not row:
        return render_template(
            "index.html",
            error="Generation not found.",
            tailored_resume=None,
            match_score=None,
            missing_keywords=[],
            recent_generations=fetch_recent_generations(),
        )

    payload_path = os.path.join(CACHE_DIR, f"{row['payload_id']}.json")
    if not os.path.exists(payload_path):
        return render_template(
            "index.html",
            error="Stored resume file not found.",
            tailored_resume=None,
            match_score=None,
            missing_keywords=[],
            recent_generations=fetch_recent_generations(),
        )

    with open(payload_path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)

    session["latest_resume_payload_id"] = row["payload_id"]

    return render_template(
        "index.html",
        error=None,
        tailored_resume=payload.get("tailored_resume"),
        match_score=payload.get("match_score", 0),
        missing_keywords=payload.get("missing_keywords", []),
        recent_generations=fetch_recent_generations(),
    )


@app.route("/modify", methods=["POST"])
def modify_resume():
    modification_text = request.form.get("modification_text", "").strip()
    if not modification_text:
        return render_template(
            "index.html",
            error="Please enter modification details.",
            tailored_resume=None,
            match_score=None,
            missing_keywords=[],
            recent_generations=fetch_recent_generations(),
        )

    payload_id = session.get("latest_resume_payload_id")
    if not payload_id:
        return render_template(
            "index.html",
            error="Generate or load a resume first.",
            tailored_resume=None,
            match_score=None,
            missing_keywords=[],
            recent_generations=fetch_recent_generations(),
        )

    payload_path = os.path.join(CACHE_DIR, f"{payload_id}.json")
    if not os.path.exists(payload_path):
        return render_template(
            "index.html",
            error="Stored resume not found. Generate again.",
            tailored_resume=None,
            match_score=None,
            missing_keywords=[],
            recent_generations=fetch_recent_generations(),
        )

    with open(payload_path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)

    base_resume = payload.get("tailored_resume", "")
    keywords = payload.get("keywords", [])
    updated_resume, merged_keywords = apply_user_modifications(base_resume, modification_text, keywords)
    score, missing = estimate_match_score(updated_resume, merged_keywords)

    structured_resume = {
        "name": infer_candidate_name(updated_resume),
        "contact": infer_contact_line(updated_resume),
        "summary": split_resume_sections(updated_resume).get("summary", ""),
        "experience": to_bullets(split_resume_sections(updated_resume).get("experience", ""), limit=50),
        "projects": to_bullets(split_resume_sections(updated_resume).get("projects", ""), limit=50),
        "skills": [s.strip() for s in re.split(r",|\n|\|", split_resume_sections(updated_resume).get("skills", "")) if s.strip()],
        "education": to_bullets(split_resume_sections(updated_resume).get("education", ""), limit=30),
    }

    new_payload_id = save_resume_payload(updated_resume, structured_resume)
    update_resume_payload(new_payload_id, merged_keywords, missing, score)
    session["latest_resume_payload_id"] = new_payload_id
    save_generation_record(False, score, modification_text, updated_resume, new_payload_id)

    return render_template(
        "index.html",
        error=None,
        tailored_resume=updated_resume,
        match_score=score,
        missing_keywords=missing,
        recent_generations=fetch_recent_generations(),
    )


@app.route("/chat", methods=["POST"])
def resume_chat():
    data = request.get_json(silent=True) or {}
    user_message = (data.get("message") or "").strip()
    if not user_message:
        return jsonify({"error": "Please enter a chat message."}), 400

    payload_id = session.get("latest_resume_payload_id")
    if not payload_id:
        return jsonify({"error": "Generate or load a resume before using chat."}), 400

    payload_path = os.path.join(CACHE_DIR, f"{payload_id}.json")
    if not os.path.exists(payload_path):
        return jsonify({"error": "Resume session expired. Generate again."}), 400

    if OpenAI is None:
        return jsonify({"error": "OpenAI SDK not installed. Run pip install -r requirements.txt"}), 500

    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        return jsonify({"error": "Set OPENAI_API_KEY environment variable."}), 500

    with open(payload_path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)

    context = build_chat_context(payload)
    client = OpenAI(api_key=api_key)

    system_prompt = (
        "You are an expert ATS resume assistant. "
        "Give concise, practical, truthful recommendations. "
        "Never invent fake achievements or experience. "
        "When user asks for rewrite, provide copy-paste ready bullets."
    )
    user_prompt = (
        f"Resume Text:\n{context['resume_text']}\n\n"
        f"Summary: {context['summary']}\n"
        f"Experience: {json.dumps(context['experience'])}\n"
        f"Projects: {json.dumps(context['projects'])}\n"
        f"Skills: {json.dumps(context['skills'])}\n"
        f"Education: {json.dumps(context['education'])}\n\n"
        f"User Message: {user_message}"
    )

    try:
        response = client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        reply = getattr(response, "output_text", "").strip()
        if not reply:
            reply = "I could not generate a response. Please try again with a more specific request."
        session["last_chat_reply"] = reply
        return jsonify({"reply": reply})
    except Exception as exc:
        return jsonify({"error": f"Chat request failed: {exc}"}), 500


@app.route("/chat/apply", methods=["POST"])
def apply_chat_suggestion():
    payload_id = session.get("latest_resume_payload_id")
    if not payload_id:
        return jsonify({"error": "Generate or load a resume before applying suggestions."}), 400

    payload_path = os.path.join(CACHE_DIR, f"{payload_id}.json")
    if not os.path.exists(payload_path):
        return jsonify({"error": "Resume session expired. Generate again."}), 400

    data = request.get_json(silent=True) or {}
    suggestion = (data.get("suggestion") or session.get("last_chat_reply") or "").strip()
    if not suggestion:
        return jsonify({"error": "No AI suggestion available to apply."}), 400

    with open(payload_path, "r", encoding="utf-8") as fh:
        payload = json.load(fh)

    base_resume = payload.get("tailored_resume", "")
    keywords = payload.get("keywords", [])
    updated_resume, merged_keywords = apply_user_modifications(base_resume, suggestion, keywords)
    score, missing = estimate_match_score(updated_resume, merged_keywords)

    sections = split_resume_sections(updated_resume)
    structured_resume = {
        "name": infer_candidate_name(updated_resume),
        "contact": infer_contact_line(updated_resume),
        "summary": sections.get("summary", ""),
        "experience": to_bullets(sections.get("experience", ""), limit=50),
        "projects": to_bullets(sections.get("projects", ""), limit=50),
        "skills": [s.strip() for s in re.split(r",|\n|\|", sections.get("skills", "")) if s.strip()],
        "education": to_bullets(sections.get("education", ""), limit=30),
    }

    new_payload_id = save_resume_payload(updated_resume, structured_resume)
    update_resume_payload(new_payload_id, merged_keywords, missing, score)
    session["latest_resume_payload_id"] = new_payload_id
    save_generation_record(False, score, "Applied AI chat suggestion", updated_resume, new_payload_id)

    return jsonify(
        {
            "tailored_resume": updated_resume,
            "match_score": score,
            "missing_keywords": missing,
            "message": "AI suggestion applied and saved as a new version.",
        }
    )


if __name__ == "__main__":
    init_db()
    port = int(os.getenv("PORT", "5000"))
    app.run(debug=True, host="0.0.0.0", port=port)
